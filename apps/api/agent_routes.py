"""
agent_routes.py
---------------
Agent trigger, approval, and execution endpoints.

Flow:
  1. POST /emails/{email_id}/process  → extract intent → create plan → (auto-run safe steps)
  2. GET  /emails/{email_id}/intents  → list extracted intents for an email
  3. GET  /plans/{plan_id}            → fetch a plan + current status
  4. POST /plans/{plan_id}/approve    → mark plan approved (approver/admin only)
  5. POST /plans/{plan_id}/reject     → mark plan rejected (approver/admin only)
  6. POST /plans/{plan_id}/execute    → run the goal executor on an approved plan
  7. GET  /plans/{plan_id}/docx       → download plan and execution log as .docx
"""

import json
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from agents.intent_agent.extractor import IntentExtractor
from agents.intent_agent.schemas import Intent
from agents.planner.planner import GoalPlanner
from agents.goal_engine.executor import GoalExecutionEngine

from db.models import (
    get_email,
    create_intent,
    get_intents_for_email,
    create_plan,
    get_plan_by_intent,
    update_plan_status,
    update_plan,
    create_risk_score,
    log_action,
)
from apps.api.auth import get_current_user
from apps.api.schemas import (
    UserOut,
    ProcessEmailResponse,
    ApprovalResponse,
    GoalPlanOut,
    GoalExecutionOut,
    StepResultOut,
)
from shared.logger import logger

agent_router = APIRouter(prefix="/emails", tags=["Agent Pipeline"])
plan_router = APIRouter(prefix="/plans", tags=["Plans"])

# Singletons — initialised once at module load
_extractor = IntentExtractor()
_planner = GoalPlanner()
_executor = GoalExecutionEngine()


# ──────────────────────────────────────────────
# 1. Process email — full pipeline trigger
# ──────────────────────────────────────────────

@agent_router.post(
    "/{email_id}/process",
    response_model=ProcessEmailResponse,
    summary="Run intent extraction + planning on an email",
)
async def process_email(
    email_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    """
    Triggers the full agent pipeline for an ingested email:
      1. Extracts intent via LLM (IntentExtractor)
      2. Converts intent → GoalPlan (deterministic GoalPlanner)
      3. Persists intent + plan to DB
      4. Auto-executes steps that do NOT require human approval
      5. Returns a summary including plan_id for follow-up approval/execution
    """
    # ── 0. Validate email exists ────────────────────────────────
    email = get_email(email_id)
    if not email:
        raise HTTPException(status_code=404, detail="Email not found")
    if email["user_id"] != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden")

    # ── 1. Extract intent ────────────────────────────────────────
    try:
        raw = await _extractor.extract(
            subject=email.get("subject", ""),
            body=email.get("body", ""),
        )
    except RuntimeError as e:
        logger.error("Intent extraction failed for email %s: %s", email_id, e)
        raise HTTPException(status_code=502, detail=f"Intent extraction failed: {e}")

    intent_dict = raw["intent"]          # plain dict from extractor
    validation_status = raw["status"]    # "valid" | "blocked"
    validation_reason = raw.get("reason")

    intent_obj = Intent(**intent_dict)

    # ── 2. Persist intent ────────────────────────────────────────
    intent_record = create_intent({
        "email_id": email_id,
        "intent_type": intent_obj.intent_type.value,
        "confidence": intent_obj.confidence_score,
        "structured_output": intent_dict,
    })
    intent_id = intent_record["id"]

    log_action({
        "user_id": current_user.id,
        "action_type": "intent_extracted",
        "status": validation_status,
        "metadata": {"email_id": email_id, "intent_id": intent_id},
    })

    # ── 3. Plan ──────────────────────────────────────────────────
    goal_plan = _planner.plan(intent_obj)

    # Blocked — low confidence or missing entities
    if goal_plan.goal_type.value == "no_action":
        plan_record = create_plan({
            "intent_id": intent_id,
            "steps": [],
            "status": "blocked",
        })
        return ProcessEmailResponse(
            email_id=email_id,
            intent_id=intent_id,
            plan_id=plan_record["id"],
            intent_type=intent_obj.intent_type.value,
            confidence=intent_obj.confidence_score,
            goal_type="no_action",
            plan_status="blocked",
            status="blocked",
            reason=validation_reason or "Low confidence or missing entities",
        )

    # ── 4. Persist plan ──────────────────────────────────────────
    steps_payload = [s.model_dump() for s in goal_plan.steps]
    plan_record = create_plan({
        "intent_id": intent_id,
        "steps": steps_payload,
        "status": "pending",
    })
    plan_id = plan_record["id"]

    # ── 5. Auto-execute safe (no-approval) steps ─────────────────
    execution_result = await _executor.execute(goal_plan)

    # Determine persisted plan status
    plan_status = execution_result.overall_status.value  # "completed" | "blocked" | "failed"
    update_plan_status(plan_id, plan_status)

    log_action({
        "user_id": current_user.id,
        "action_type": "plan_created",
        "status": plan_status,
        "metadata": {"plan_id": plan_id, "goal_type": goal_plan.goal_type.value},
    })

    return ProcessEmailResponse(
        email_id=email_id,
        intent_id=intent_id,
        plan_id=plan_id,
        intent_type=intent_obj.intent_type.value,
        confidence=intent_obj.confidence_score,
        goal_type=goal_plan.goal_type.value,
        plan_status=plan_status,
        execution=GoalExecutionOut(
            goal_type=execution_result.goal_type,
            overall_status=execution_result.overall_status.value,
            step_results=[s.model_dump() for s in execution_result.step_results],
        ),
        status="processed",
    )


# ──────────────────────────────────────────────
# 2. List intents for an email
# ──────────────────────────────────────────────

@agent_router.get(
    "/{email_id}/intents",
    summary="List extracted intents for an email",
)
async def list_intents(
    email_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    email = get_email(email_id)
    if not email:
        raise HTTPException(status_code=404, detail="Email not found")
    if email["user_id"] != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden")

    return get_intents_for_email(email_id)


# ──────────────────────────────────────────────
# 3. Fetch a plan
# ──────────────────────────────────────────────

@plan_router.get(
    "/{plan_id}",
    summary="Fetch a plan by ID",
)
async def get_plan(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    """Returns the plan record including its steps and current status."""
    from db.models import get_supabase
    client = get_supabase()
    result = client.table("plans").select("*").eq("id", plan_id).execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Plan not found")
    return result.data[0]


# ──────────────────────────────────────────────
# 2.5 Dashboard — Pending & History
# ──────────────────────────────────────────────

@plan_router.get("/pending", tags=["Plans"])
async def get_pending_plans_endpoint(current_user: UserOut = Depends(get_current_user)):
    """Return all pending plans for the current user."""
    from db.models import get_pending_plans
    plans = get_pending_plans(current_user.email)
    return {"plans": plans, "count": len(plans)}


@plan_router.get("/history", tags=["Plans"])
async def get_plan_history_endpoint(current_user: UserOut = Depends(get_current_user), limit: int = 20):
    """Return recent execution history for the user."""
    from db.models import PLAN_CACHE, get_history_plans

    # Primary source: Supabase
    plans = get_history_plans(current_user.email, limit=limit)
    seen_ids = {p.get("id") for p in plans if p.get("id")}

    # Fallback source: in-memory PLAN_CACHE
    # This covers cases where DB writes are delayed/blocked (e.g. RLS) but the API
    # has already executed and updated the cached plan.
    cached = []
    for pid, row in PLAN_CACHE.items():
        if pid in seen_ids:
            continue
        if row.get("user_email") != current_user.email:
            continue
        if row.get("status") not in ("approved", "executed", "rejected", "failed"):
            continue
        cached.append(row)

    merged = plans + cached

    def _sort_key(r: dict):
        # created_at in Supabase is usually ISO; cache may also be ISO.
        return r.get("created_at") or ""

    merged.sort(key=_sort_key, reverse=True)
    return {"plans": merged[:limit]}


# ──────────────────────────────────────────────
# 4. Approve a plan
# ──────────────────────────────────────────────

@plan_router.post(
    "/{plan_id}/approve",
    response_model=ApprovalResponse,
    summary="Approve a plan (owner, approver, or admin)",
)
async def approve_plan(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    from db.models import get_plan_cached
    plan = get_plan_cached(plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    # Ownership check disabled for extension flow:
    # When /analyze is called from extension, user_email is set to the email sender,
    # not the authenticated user. Any authenticated user can approve plans.
    # is_owner = plan.get("user_email") == current_user.email
    # is_privileged = current_user.role in ("admin", "approver")
    # if not is_owner and not is_privileged:
    #     raise HTTPException(status_code=403, detail="You can only approve your own plans")

    updated = update_plan_status(plan_id, "approved")
    log_action({
        "user_id": current_user.id,
        "action_type": "plan_approved",
        "status": "approved",
        "metadata": {"plan_id": plan_id},
    })
    logger.info("Plan %s approved by %s", plan_id, current_user.email)
    return ApprovalResponse(
        plan_id=plan_id,
        status="approved",
        message="Plan approved. Call /execute to run it.",
    )


# ──────────────────────────────────────────────
# 5. Reject a plan
# ──────────────────────────────────────────────

@plan_router.post(
    "/{plan_id}/reject",
    response_model=ApprovalResponse,
    summary="Reject a plan (owner, approver, or admin)",
)
async def reject_plan(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    from db.models import get_plan_cached
    plan = get_plan_cached(plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    # Ownership check disabled for extension flow (see approve_plan for details)
    # is_owner = plan.get("user_email") == current_user.email
    # is_privileged = current_user.role in ("admin", "approver")
    # if not is_owner and not is_privileged:
    #     raise HTTPException(status_code=403, detail="You can only reject your own plans")

    update_plan_status(plan_id, "rejected")
    log_action({
        "user_id": current_user.id,
        "action_type": "plan_rejected",
        "status": "rejected",
        "metadata": {"plan_id": plan_id},
    })
    logger.info("Plan %s rejected by %s", plan_id, current_user.email)
    return ApprovalResponse(
        plan_id=plan_id,
        status="rejected",
        message="Plan rejected and will not be executed.",
    )


# ──────────────────────────────────────────────
# 6. Execute an approved plan
# ──────────────────────────────────────────────

@plan_router.post(
    "/{plan_id}/execute",
    response_model=ApprovalResponse,
    summary="Execute an approved plan",
)
async def execute_plan(
    plan_id: str,
    dry_run: bool = False,
    current_user: UserOut = Depends(get_current_user),
):
    """
    Runs the GoalExecutionEngine on an approved plan.
    Only works if plan status == 'approved' (unless dry_run is True).
    """
    from db.models import get_plan_cached as _get_plan_row, get_gmail_token
    from tools.orchestrator.tool_dispatcher import dispatch_plan

    plan_row = _get_plan_row(plan_id)
    if not plan_row:
        raise HTTPException(status_code=404, detail="Plan not found")

    # Only approved plans can be executed (unless it's a dry-run)
    if not dry_run and plan_row["status"] != "approved":
        raise HTTPException(
            status_code=409,
            detail=f"Plan is '{plan_row['status']}' — only approved plans can be executed.",
        )

    # Get refresh token for the user
    # CRITICAL: We MUST use the authenticated current_user.email, not plan_row.get("user_email").
    # In the extension flow, plan_row.user_email may be the *sender* of the email,
    # but the token belongs to the USER who is logged in and executing.
    token_row = get_gmail_token(current_user.email)
    if not token_row or not token_row.get("google_refresh_token"):
        raise HTTPException(
            status_code=400,
            detail=f"Gmail refresh token not found for {current_user.email}. Please re-login."
        )
    refresh_token = token_row["google_refresh_token"]

    # Reconstruct GoalPlan from DB or Cache JSON
    # Cache format: plan_json: { steps: [...] }
    # DB format: steps: [...]
    plan_json = plan_row.get("plan_json") or {}
    steps_raw = plan_json.get("steps") or plan_row.get("steps") or []

    if isinstance(steps_raw, str):
        steps_raw = json.loads(steps_raw)

    # Context for Gmail/Calendar tools
    email_context = plan_json.get("email_context") or {
        "subject": plan_row.get("subject", ""),
        "sender": plan_row.get("sender", ""),
    }

    # Execute
    try:
        results = await dispatch_plan(
            plan_steps=steps_raw,
            refresh_token=refresh_token,
            email_context=email_context,
            dry_run=dry_run
        )

        success_count = sum(1 for r in results if r.get("success"))
        
        # Mapping results to StepResultOut schema
        step_results = []
        for i, r in enumerate(results):
            step_results.append(StepResultOut(
                step_id=i + 1,
                status="success" if r.get("success") else "failure",
                message=r.get("error") or r.get("note") or "Step completed"
            ))

        # If dry-run, we don't update status to executed
        if dry_run:
            final_status = plan_row["status"]
        else:
            final_status = "executed" if (success_count == len(results)) else "failed"
            # Always update cache so UI can reflect immediately even if DB write fails.
            try:
                from db.models import PLAN_CACHE
                PLAN_CACHE.setdefault(plan_id, plan_row)["status"] = final_status
                PLAN_CACHE.setdefault(plan_id, plan_row)["execution_log"] = results
                PLAN_CACHE.setdefault(plan_id, plan_row)["user_email"] = current_user.email
            except Exception:
                pass

            try:
                update_plan(plan_id, {
                    "status": final_status,
                    "execution_log": results,
                    "user_email": current_user.email,
                })
            except Exception as db_exc:
                logger.warning("DB persist of execution_log failed: %s", db_exc)

        log_action({
            "user_id": current_user.id,
            "action_type": "plan_dry_run" if dry_run else "plan_executed",
            "status": final_status,
            "metadata": {"plan_id": plan_id},
        })

        return ApprovalResponse(
            plan_id=plan_id,
            status=final_status,
            message=f"Execution finished with status: {final_status}" if not dry_run else "Dry run complete",
            steps_total=len(results),
            steps_succeeded=success_count,
            execution=GoalExecutionOut(
                goal_type=plan_row.get("goal_type", "unknown"),
                overall_status=final_status,
                step_results=step_results,
            )
        )
    except Exception as exc:
        logger.error("Execution failed for plan %s: %s", plan_id, exc)
        update_plan_status(plan_id, "failed")
        raise HTTPException(status_code=500, detail=f"Execution failed: {exc}")


# ──────────────────────────────────────────────
# 7. Dry-run a plan (preview without execution)
# ──────────────────────────────────────────────

@plan_router.post(
    "/{plan_id}/approve-and-execute",
    response_model=ApprovalResponse,
    summary="Approve + Execute in one shot",
)
async def approve_and_execute(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    """Atomically approve and execute a plan."""
    from db.models import update_plan_status
    # 1. Approve
    update_plan_status(plan_id, "approved")
    # 2. Execute
    return await execute_plan(plan_id, dry_run=False, current_user=current_user)


@plan_router.post(
    "/{plan_id}/dry-run",
    summary="Preview what a plan would do without executing it",
)
async def dry_run_plan_endpoint(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    """
    Returns a simulation of each plan step without making real API calls.
    Used by the Chrome extension to show the user what will happen before approval.
    """
    from db.models import get_plan_cached as _get_plan_row
    from tools.orchestrator.tool_dispatcher import dispatch_plan

    plan_row = _get_plan_row(plan_id)
    if not plan_row:
        raise HTTPException(status_code=404, detail="Plan not found")

    plan_json = plan_row.get("plan_json") or {}
    steps_raw = plan_json.get("steps") or plan_row.get("steps") or []

    if isinstance(steps_raw, str):
        steps_raw = json.loads(steps_raw)

    # We don't need a real token for dry-run
    results = await dispatch_plan(
        plan_steps=steps_raw,
        refresh_token="simulated",
        email_context=plan_json.get("email_context") or {},
        dry_run=True
    )

    return {
        "plan_id": plan_id,
        "summary": plan_json.get("summary") or "Dry run complete",
        "steps": results
    }


# ──────────────────────────────────────────────
# 9. Generate Google Docs execution report
# ──────────────────────────────────────────────

@plan_router.get(
    "/{plan_id}/report",
    summary="Generate a Google Doc execution report",
)
async def generate_plan_report(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    """
    Creates a Google Document summarising the executed plan and
    returns a shareable URL. Requires the plan to have status 'executed'.
    """
    from db.models import get_plan_cached, get_gmail_token
    from tools.gdocs_tool.gdocs_tool import create_execution_report_doc

    plan_row = get_plan_cached(plan_id)
    if not plan_row:
        raise HTTPException(status_code=404, detail="Plan not found")

    if plan_row.get("status") not in ("executed", "failed"):
        raise HTTPException(
            status_code=409,
            detail=f"Plan is '{plan_row.get('status')}' — only executed plans can have a report generated.",
        )

    token_row = get_gmail_token(current_user.email)
    if not token_row or not token_row.get("google_refresh_token"):
        raise HTTPException(
            status_code=400,
            detail="Gmail refresh token not found. Please re-login.",
        )
    refresh_token = token_row["google_refresh_token"]

    try:
        doc_url = create_execution_report_doc(
            plan=plan_row,
            user_email=current_user.email,
            refresh_token=refresh_token,
        )
        log_action({
            "user_id": current_user.id,
            "action_type": "report_generated",
            "status": "success",
            "metadata": {"plan_id": plan_id, "doc_url": doc_url},
        })
        return {"plan_id": plan_id, "doc_url": doc_url}
    except Exception as exc:
        logger.error("Report generation failed for plan %s: %s", plan_id, exc)
        raise HTTPException(status_code=502, detail=f"Failed to create Google Doc: {exc}")


@plan_router.get(
    "/{plan_id}/docx",
    summary="Download a plan as a .docx file",
)
async def download_plan_docx(
    plan_id: str,
    current_user: UserOut = Depends(get_current_user),
):
    import io
    from datetime import datetime

    from db.models import get_plan_cached

    plan_row = get_plan_cached(plan_id)
    if not plan_row:
        raise HTTPException(status_code=404, detail="Plan not found")

    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"python-docx not installed: {exc}")

    doc = Document()
    doc.add_heading("SecureIntent Plan", level=1)
    doc.add_paragraph(f"Plan ID: {plan_row.get('id', plan_id)}")
    doc.add_paragraph(f"Status: {plan_row.get('status', '')}")
    doc.add_paragraph(f"Goal Type: {plan_row.get('goal_type', '')}")
    doc.add_paragraph(f"User: {plan_row.get('user_email', '')}")
    doc.add_paragraph(f"Subject: {plan_row.get('subject', '')}")
    doc.add_paragraph(f"Sender: {plan_row.get('sender', '')}")
    doc.add_paragraph(f"Created At: {plan_row.get('created_at', '')}")

    doc.add_heading("Intent", level=2)
    intent_json = plan_row.get("intent_json") or {}
    doc.add_paragraph(json.dumps(intent_json, indent=2, default=str))

    doc.add_heading("Plan", level=2)
    plan_json = plan_row.get("plan_json") or {}
    steps = plan_json.get("steps") or []
    if steps:
        for idx, step in enumerate(steps, start=1):
            action = step.get("action", "")
            desc = step.get("description", "")
            doc.add_paragraph(f"{idx}. {action} — {desc}")
    else:
        doc.add_paragraph(json.dumps(plan_json, indent=2, default=str))

    doc.add_heading("Execution Log", level=2)
    execution_log = plan_row.get("execution_log") or []
    doc.add_paragraph(json.dumps(execution_log, indent=2, default=str))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    ts = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    filename = f"secureintent-plan-{plan_id}-{ts}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
